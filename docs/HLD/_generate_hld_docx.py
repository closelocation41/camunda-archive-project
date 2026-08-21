from __future__ import annotations

import html
import math
import re
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
OUT = DOCS / "HLD"
DIAGRAMS = OUT / "diagrams"

SOURCES = [
    ("WHITE_PAPER.md", "Camunda Data Archiving Project White Paper", "Camunda_Data_Archiving_White_Paper.docx"),
    ("ARCHITECTURE.md", "Camunda Archive HLD - Architecture", "Camunda_Archive_HLD_Architecture.docx"),
    ("camunda-database-docs.md", "Camunda Archive HLD - Database Reference", "Camunda_Archive_HLD_Database_Reference.docx"),
    ("camunda-archive-session-6.md", "Camunda Archive Session 6 - History Tables, Byte Array Dependencies, and Manual Archive Plan", "Camunda_Archive_Session_6_History_Tables_ByteArray_Archival.docx"),
    ("camunda-history-cleanup.md", "Camunda Archive HLD - History Cleanup", "Camunda_Archive_HLD_History_Cleanup.docx"),
    ("RESTORE_DESIGN.md", "Camunda Archive HLD - Restore Design", "Camunda_Archive_HLD_Restore_Design.docx"),
    ("API.md", "Camunda Archive HLD - API Reference", "Camunda_Archive_HLD_API_Reference.docx"),
    ("OPERATIONS.md", "Camunda Archive HLD - Operations Guide", "Camunda_Archive_HLD_Operations_Guide.docx"),
]


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 99, 111)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
INK = "0B2545"
BORDER = "C8D1DD"


def font(name="Calibri", size=11, color=None, bold=None, italic=None):
    return {"name": name, "size": size, "color": color, "bold": bold, "italic": italic}


def set_run(run, spec):
    run.font.name = spec.get("name", "Calibri")
    run._element.rPr.rFonts.set(qn("w:ascii"), run.font.name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), run.font.name)
    run.font.size = Pt(spec.get("size", 11))
    if spec.get("color"):
        run.font.color.rgb = spec["color"]
    if spec.get("bold") is not None:
        run.bold = spec["bold"]
    if spec.get("italic") is not None:
        run.italic = spec["italic"]


def paragraph_border_bottom(paragraph, color="D7DBE2", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=RGBColor(0, 0, 0), size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_run(r, font(size=size, color=color, bold=bold))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("Camunda Data Archiving Project")
    set_run(r, font(size=9, color=GRAY, bold=True))
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    r = footer.add_run("HLD Documentation Set")
    set_run(r, font(size=9, color=GRAY))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, font(size=24, color=NAVY, bold=True))
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("High-Level Design and Technical Reference")
    set_run(r, font(size=12, color=GRAY, italic=True))
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, "9EB7D5")


def add_code_block(doc, code, label="Code"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run(r, font(size=9, color=DARK_BLUE, bold=True))
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FB")
    cell.text = ""
    for idx, line in enumerate(code.strip("\n").splitlines() or [""]):
        if idx:
            cell.add_paragraph()
        para = cell.paragraphs[-1]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        set_run(run, {"name": "Consolas", "size": 8.5, "color": RGBColor(30, 30, 30)})


def wrap_text(draw, text, font_obj, width):
    words = str(text).split()
    lines = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= width or not line:
            line = test
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines[:4]


def load_font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_node(draw, box, text, fill="#F4F6F9", outline="#6B8BB5"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = load_font(22, True)
    lines = wrap_text(draw, text, title_font, x2 - x1 - 28)
    total = len(lines) * 27
    y = y1 + ((y2 - y1 - total) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#0B2545", font=title_font)
        y += 27


def arrow(draw, start, end, color="#4E6D8F"):
    draw.line([start, end], fill=color, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 14
    points = [
        end,
        (end[0] - head * math.cos(angle - 0.5), end[1] - head * math.sin(angle - 0.5)),
        (end[0] - head * math.cos(angle + 0.5), end[1] - head * math.sin(angle + 0.5)),
    ]
    draw.polygon(points, fill=color)


def parse_label(raw, fallback):
    m = re.search(r"\[([^\]]+)\]|\(([^)]+)\)|\{([^}]+)\}", raw)
    if m:
        return html.unescape(next(g for g in m.groups() if g))
    return fallback


def remember_label(labels, node_id, raw):
    parsed = parse_label(raw, node_id)
    if parsed != node_id or node_id not in labels:
        labels[node_id] = parsed


def flowchart_image(code, path, title):
    direction = "LR" if re.search(r"flowchart\s+LR", code) else "TB"
    labels = {}
    edges = []
    for line in code.splitlines():
        line = line.strip()
        if not line or line.startswith("flowchart"):
            continue
        parts = re.split(r"\s*--?>\s*|\s*-->>\s*", line)
        if len(parts) >= 2:
            left, right = parts[0].strip(), parts[-1].strip()
            left_id = re.match(r"([A-Za-z0-9_]+)", left)
            right_id = re.match(r"([A-Za-z0-9_]+)", right)
            if left_id and right_id:
                l_id, r_id = left_id.group(1), right_id.group(1)
                remember_label(labels, l_id, left)
                remember_label(labels, r_id, right)
                edges.append((l_id, r_id))
        else:
            m = re.match(r"([A-Za-z0-9_]+)(.*)", line)
            if m:
                remember_label(labels, m.group(1), line)
    nodes = list(labels.keys())[:12]
    if not nodes:
        return False
    if direction == "LR":
        w, h = 1800, max(700, 230 + 140 * math.ceil(len(nodes) / 4))
        cols = min(4, len(nodes))
        rows = math.ceil(len(nodes) / cols)
        positions = {}
        for i, node in enumerate(nodes):
            col = i % cols
            row = i // cols
            x = 90 + col * ((w - 180) / cols)
            y = 150 + row * 150
            positions[node] = (x, y, x + 310, y + 92)
    else:
        w, h = 1400, max(800, 170 + len(nodes) * 130)
        positions = {}
        for i, node in enumerate(nodes):
            x = 485
            y = 120 + i * 120
            positions[node] = (x, y, x + 430, y + 82)
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    tfont = load_font(30, True)
    draw.text((60, 35), title, fill="#0B2545", font=tfont)
    for left, right in edges:
        if left in positions and right in positions:
            a = positions[left]
            b = positions[right]
            arrow(draw, ((a[0] + a[2]) / 2, a[3]), ((b[0] + b[2]) / 2, b[1]))
    for i, node in enumerate(nodes):
        fill = "#E8EEF5" if i % 2 == 0 else "#F4F6F9"
        draw_node(draw, positions[node], labels[node], fill=fill)
    img.save(path)
    return True


def sequence_image(code, path, title):
    participants = []
    messages = []
    for line in code.splitlines():
        line = line.strip()
        m = re.match(r"(participant|actor)\s+(\w+)(?:\s+as\s+(.+))?", line)
        if m:
            participants.append((m.group(2), m.group(3) or m.group(2)))
        m = re.match(r"(\w+)\s*[-=]+>>\s*(\w+)\s*:\s*(.+)", line)
        if m:
            messages.append(m.groups())
            for p in [m.group(1), m.group(2)]:
                if not any(existing[0] == p for existing in participants):
                    participants.append((p, p))
    participants = participants[:7]
    messages = messages[:14]
    if not participants:
        return False
    w = max(1400, 200 + len(participants) * 210)
    h = 230 + len(messages) * 80
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    tfont = load_font(30, True)
    small = load_font(19, False)
    bold = load_font(19, True)
    draw.text((60, 35), title, fill="#0B2545", font=tfont)
    xs = {p[0]: 100 + i * ((w - 200) / max(1, len(participants) - 1)) for i, p in enumerate(participants)}
    for pid, label in participants:
        x = xs[pid]
        draw.rounded_rectangle((x - 75, 100, x + 75, 148), radius=12, fill="#E8EEF5", outline="#6B8BB5", width=2)
        lines = wrap_text(draw, label, bold, 130)
        draw.text((x - 65, 114), lines[0], fill="#0B2545", font=bold)
        draw.line((x, 148, x, h - 45), fill="#C8D1DD", width=2)
    y = 190
    for src, dst, msg in messages:
        if src not in xs or dst not in xs:
            continue
        x1, x2 = xs[src], xs[dst]
        arrow(draw, (x1, y), (x2, y))
        label = " ".join(wrap_text(draw, msg, small, abs(x2 - x1) + 170))
        draw.text((min(x1, x2) + 12, y - 26), label[:90], fill="#333333", font=small)
        y += 80
    img.save(path)
    return True


def er_image(code, path, title):
    entities = []
    rels = []
    for line in code.splitlines():
        line = line.strip()
        m = re.match(r"([A-Z0-9_]+)\s+.+?\s+([A-Z0-9_]+)\s*:", line)
        if m:
            rels.append(m.groups())
            for entity in m.groups():
                if entity not in entities:
                    entities.append(entity)
    entities = entities[:14]
    if not entities:
        return False
    w, h = 1600, 1000
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    tfont = load_font(30, True)
    bold = load_font(19, True)
    draw.text((60, 35), title, fill="#0B2545", font=tfont)
    cols = 3
    positions = {}
    for i, entity in enumerate(entities):
        col = i % cols
        row = i // cols
        x = 95 + col * 500
        y = 120 + row * 160
        positions[entity] = (x, y, x + 360, y + 82)
    for a, b in rels:
        if a in positions and b in positions:
            pa, pb = positions[a], positions[b]
            arrow(draw, ((pa[0] + pa[2]) / 2, (pa[1] + pa[3]) / 2), ((pb[0] + pb[2]) / 2, (pb[1] + pb[3]) / 2), "#AAB7C6")
    for entity, box in positions.items():
        draw_node(draw, box, entity.replace("_", " "), fill="#F4F6F9")
    img.save(path)
    return True


def render_mermaid(code, path, title):
    first = code.strip().splitlines()[0] if code.strip() else ""
    if first.startswith("sequenceDiagram"):
        return sequence_image(code, path, title)
    if first.startswith("erDiagram"):
        return er_image(code, path, title)
    if first.startswith("flowchart") or first.startswith("graph"):
        return flowchart_image(code, path, title)
    return False


def split_markdown(md):
    pattern = re.compile(r"```(\w+)?\n(.*?)```", re.DOTALL)
    pos = 0
    for match in pattern.finditer(md):
        if match.start() > pos:
            yield ("text", None, md[pos : match.start()])
        yield ("code", match.group(1) or "", match.group(2))
        pos = match.end()
    if pos < len(md):
        yield ("text", None, md[pos:])


def inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(([^)]+)\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            r = paragraph.add_run(text[pos : match.start()])
            set_run(r, font())
        token = match.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run(r, font(bold=True))
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_run(r, {"name": "Consolas", "size": 9.5, "color": RGBColor(30, 30, 30)})
        elif token.startswith("["):
            label = token[1 : token.find("]")]
            url = match.group(2)
            add_hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run(r, font())


def add_markdown_table(doc, rows):
    cleaned = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if cells and not all(re.match(r"^:?-{3,}:?$", c) for c in cells):
            cleaned.append(cells)
    if len(cleaned) < 1:
        return
    max_cols = max(len(r) for r in cleaned)
    table = doc.add_table(rows=len(cleaned), cols=max_cols)
    table.style = "Table Grid"
    base_width = 6.5 / max_cols
    widths = [base_width] * max_cols
    set_table_width(table, widths)
    for r_idx, row in enumerate(cleaned):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, LIGHT_GRAY)
                set_cell_text(cell, text, bold=True, color=NAVY, size=8.6)
            else:
                set_cell_text(cell, text, size=8.4)


def add_text_block(doc, text):
    table_buf = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line)
            continue
        if table_buf:
            add_markdown_table(doc, table_buf)
            table_buf = []
        if not line.strip():
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            inline_runs(p, re.sub(r"^\d+\.\s+", "", line).strip())
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            inline_runs(p, re.sub(r"^[-*]\s+", "", line).strip())
        else:
            p = doc.add_paragraph()
            inline_runs(p, line.strip())
    if table_buf:
        add_markdown_table(doc, table_buf)


def markdown_to_docx(source_name, title, output_name):
    md = (DOCS / source_name).read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc, title)
    diagram_count = 0
    figure_count = 0
    for kind, lang, content in split_markdown(md):
        if kind == "text":
            add_text_block(doc, content)
        elif lang == "mermaid":
            figure_count += 1
            path = DIAGRAMS / f"{Path(output_name).stem}_figure_{figure_count}.png"
            ok = render_mermaid(content, path, f"Figure {figure_count}")
            if ok:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(path), width=Inches(6.3))
                cap = doc.add_paragraph(f"Figure {figure_count}: Diagram generated from the source Markdown.")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
                diagram_count += 1
            else:
                add_code_block(doc, content, "Mermaid diagram source")
        else:
            add_code_block(doc, content, f"{lang or 'code'} block")
    doc.core_properties.title = title
    doc.core_properties.subject = "Camunda archive high-level design documentation"
    doc.core_properties.keywords = "Camunda, archive, restore, history cleanup, HLD, white paper"
    out_path = OUT / output_name
    doc.save(out_path)
    return out_path, diagram_count


def create_index(results):
    doc = Document()
    configure_document(doc, "Camunda Archive HLD Documentation Index")
    p = doc.add_paragraph()
    inline_runs(
        p,
        "This index links the generated DOCX documents for the Camunda Data Archiving high-level design package. "
        "The Markdown source files remain in the parent docs folder.",
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [1.6, 2.1, 1.0, 1.8])
    headers = ["Document", "Purpose", "Diagrams", "Link"]
    for idx, h in enumerate(headers):
        shade_cell(table.cell(0, idx), LIGHT_GRAY)
        set_cell_text(table.cell(0, idx), h, bold=True, color=NAVY)
    purposes = {
        "Camunda_Data_Archiving_White_Paper.docx": "Primary enterprise white paper and official technical reference.",
        "Camunda_Archive_HLD_Architecture.docx": "Architecture overview and current approach summary.",
        "Camunda_Archive_HLD_Database_Reference.docx": "Camunda table analysis and archive table mappings.",
        "Camunda_Archive_Session_6_History_Tables_ByteArray_Archival.docx": "Session 6 archive plan covering history tables, byte-array dependencies, and the manual archive workflow.",
        "Camunda_Archive_HLD_History_Cleanup.docx": "History Cleanup lifecycle and integration guidance.",
        "Camunda_Archive_HLD_Restore_Design.docx": "Restore/re-sync design and constraints.",
        "Camunda_Archive_HLD_API_Reference.docx": "API surface and roles.",
        "Camunda_Archive_HLD_Operations_Guide.docx": "Operator runbook and production guidance.",
    }
    for path, diagrams in results:
        row = table.add_row().cells
        set_cell_text(row[0], path.name, bold=True, color=NAVY)
        set_cell_text(row[1], purposes[path.name], size=8.2)
        set_cell_text(row[2], str(diagrams), size=8.2)
        row[3].text = ""
        p = row[3].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, "Open DOCX", path.resolve().as_uri())
    out_path = OUT / "Camunda_Archive_HLD_Document_Index.docx"
    doc.save(out_path)
    return out_path


def structural_check(path):
    with zipfile.ZipFile(path) as zf:
        required = ["[Content_Types].xml", "word/document.xml"]
        for item in required:
            if item not in zf.namelist():
                raise RuntimeError(f"{path.name} is missing {item}")
        xml = zf.read("word/document.xml")
        if b"w:document" not in xml:
            raise RuntimeError(f"{path.name} does not look like a Word document")
    return True


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    results = []
    for source, title, output in SOURCES:
        out_path, diagrams = markdown_to_docx(source, title, output)
        structural_check(out_path)
        results.append((out_path, diagrams))
    index = create_index(results)
    structural_check(index)
    for path, diagrams in results:
        print(f"{path} | diagrams={diagrams}")
    print(index)


if __name__ == "__main__":
    main()
